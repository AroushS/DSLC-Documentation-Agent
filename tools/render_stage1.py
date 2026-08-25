"""
render_stage1.py — DSLC Stage 1 Standard Document Renderer

Convert a validated Stage 1 Markdown source into a professional editable
Microsoft Word (DOCX) document and, where locally supported, a PDF.

Usage:
    python tools/render_stage1.py stage1_draft.md
    python tools/render_stage1.py stage1_draft.md --output-dir output

Inputs:
    - Validated Stage 1 Markdown file
    - Optional: corporate branding template (templates/corporate_stage1_template.docx)

Outputs:
    - output/Stage1_<ProjectName>.docx (editable master)
    - output/Stage1_<ProjectName>.pdf (if conversion available)

Architectural boundaries:
    - Renderer controls PRESENTATION ONLY
    - Never modifies governance status
    - Never invents evidence
    - Never alters project facts
    - Never recalculates Stage 1 completion
    - Never implies approval without evidence

Exit codes:
    0 = SUCCESS (DOCX created, optional PDF created or reported unavailable)
    1 = FAILURE (input not found, validation errors, etc.)
"""

import re
import sys
import argparse
from pathlib import Path
from datetime import date
from dataclasses import dataclass, field
import subprocess
from urllib.parse import quote

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not found. Run: pip install python-docx")
    sys.exit(1)


# ============================================================================
# COLOURS (Professional corporate palette)
# ============================================================================

# Sky primary brand colours
SKY_BLUE   = RGBColor(0x00, 0x53, 0x9F)   # H1, main headers, structural dividers
SKY_PURPLE = RGBColor(0x6B, 0x2D, 0x8B)   # H2, secondary accents, DRAFT emphasis
SKY_PINK   = RGBColor(0xD4, 0x14, 0x5A)   # Sparse accent / cover-page separators
SKY_ORANGE = RGBColor(0xF0, 0x7F, 0x1E)   # PENDING, blockers, attention

# Neutral text and surface palette
TEXT_DARK       = RGBColor(0x1A, 0x1A, 0x1A)   # Primary body text
TEXT_GREY       = RGBColor(0x6B, 0x6B, 0x6B)   # Secondary / caption text
LIGHT_GREY      = RGBColor(0xCC, 0xCC, 0xCC)   # Table borders
VERY_LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)   # Alternate row shading
WHITE           = RGBColor(0xFF, 0xFF, 0xFF)    # Page / cell backgrounds

# Status indicator colours (text only; never full-cell background fill)
STATUS_GREEN = RGBColor(0x1A, 0x7A, 0x40)   # Complete / Evidence Present
STATUS_AMBER = RGBColor(0xD4, 0x6A, 0x0A)   # Partial / clarification required
STATUS_RED   = RGBColor(0xB3, 0x1B, 0x1B)   # Not ready / blocking outcome

# Hex equivalents (for python-docx XML shading operations)
HEX_SKY_BLUE        = "00539F"
HEX_SKY_PURPLE      = "6B2D8B"
HEX_SKY_ORANGE      = "F07F1E"
HEX_WHITE           = "FFFFFF"
HEX_VERY_LIGHT_GREY = "F5F5F5"
HEX_LIGHT_GREY      = "CCCCCC"
HEX_TEXT_DARK       = "1A1A1A"

# Light status fills for the executive dashboard.
# Written status always remains visible; colour is only a visual aid.
HEX_STATUS_GREEN_LIGHT = "EAF5EE"
HEX_STATUS_AMBER_LIGHT = "FFF4E5"
HEX_STATUS_RED_LIGHT   = "FDECEC"
HEX_STATUS_GREY_LIGHT  = "F2F2F2"

# Status keyword → (text RGBColor, hex string)
STATUS_COLOUR: dict[str, tuple[RGBColor, str]] = {
    "complete":       (STATUS_GREEN, "1A7A40"),
    "pending":        (STATUS_AMBER, "D46A0A"),
    "draft":          (SKY_PURPLE, "6B2D8B"),
    "not applicable": (TEXT_GREY, "6B6B6B"),
    "risk accepted":  (TEXT_GREY, "6B6B6B"),
}

# Page layout — A4 with ~2.2 cm margins
BODY_CM = 16.6


# ============================================================================
# MARKDOWN PARSER
# ============================================================================

@dataclass
class Block:
    kind: str                           # heading | paragraph | table | list | blockquote | code | rule
    level: int = 0
    text: str = ""
    items: list = field(default_factory=list)
    headers: list = field(default_factory=list)
    rows: list = field(default_factory=list)
    ordered: bool = False


def _parse_table(table_lines: list[str]) -> tuple[list, list]:
    """Parse Markdown table lines into headers and rows."""
    headers: list[str] = []
    rows: list[list[str]] = []
    for i, ln in enumerate(table_lines):
        if re.match(r"^\|[-| :]+\|", ln.strip()):
            continue
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if i == 0 or not headers:
            headers = cells
        else:
            rows.append(cells)
    return headers, rows


def parse_markdown(lines: list[str]) -> list[Block]:
    """Parse Markdown text into block structure."""
    blocks: list[Block] = []
    i = 0
    while i < len(lines):
        ln = lines[i].rstrip()

        # Fenced code block
        if ln.startswith("```"):
            code: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                code.append(lines[i].rstrip())
                i += 1
            blocks.append(Block(kind="code", text="\n".join(code)))
            i += 1
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)", ln)
        if m:
            text = re.sub(r"\s*\{#[^}]+\}$", "", m.group(2)).strip()
            blocks.append(Block(kind="heading", level=len(m.group(1)), text=text))
            i += 1
            continue

        # Table
        if ln.startswith("|"):
            tbl: list[str] = []
            while i < len(lines) and lines[i].rstrip().startswith("|"):
                tbl.append(lines[i].rstrip())
                i += 1
            h, r = _parse_table(tbl)
            blocks.append(Block(kind="table", headers=h, rows=r))
            continue

        # Horizontal rule
        if re.match(r"^(---+|\*\*\*+|___+)\s*$", ln):
            blocks.append(Block(kind="rule"))
            i += 1
            continue

        # Blockquote
        if ln.startswith(">"):
            bq: list[str] = []
            while i < len(lines) and lines[i].rstrip().startswith(">"):
                bq.append(lines[i].rstrip()[1:].strip())
                i += 1
            blocks.append(Block(kind="blockquote", text=" ".join(bq)))
            continue

        # Unordered list
        if re.match(r"^[-*+]\s+", ln):
            items: list[str] = []
            while i < len(lines):
                m2 = re.match(r"^[-*+]\s+(.*)", lines[i].rstrip())
                if m2:
                    items.append(m2.group(1))
                    i += 1
                elif lines[i].startswith("  ") and items:
                    items[-1] += " " + lines[i].strip()
                    i += 1
                else:
                    break
            blocks.append(Block(kind="list", items=items, ordered=False))
            continue

        # Ordered list
        if re.match(r"^\d+\.\s+", ln):
            items = []
            while i < len(lines):
                m2 = re.match(r"^\d+\.\s+(.*)", lines[i].rstrip())
                if m2:
                    items.append(m2.group(1))
                    i += 1
                else:
                    break
            blocks.append(Block(kind="list", items=items, ordered=True))
            continue

        # Empty line
        if not ln.strip():
            i += 1
            continue

        # Paragraph
        para: list[str] = []
        while i < len(lines):
            ll = lines[i].rstrip()
            if not ll:
                break
            if re.match(r"^(#{1,6})\s+|^\||^[-*+]\s+|^\d+\.\s+|^>|^```", ll):
                break
            if re.match(r"^(---+|\*\*\*+|___+)\s*$", ll):
                break
            para.append(ll)
            i += 1
        if para:
            blocks.append(Block(kind="paragraph", text=" ".join(para)))

    return blocks


# ============================================================================
# DOCUMENT ANALYSIS
# ============================================================================

def _strip_md(text: str) -> str:
    """Remove inline Markdown formatting."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def _status_colour(status_text: str) -> tuple[RGBColor, str] | None:
    """Look up presentation colour for an existing governance status."""
    status = status_text.strip().lower()
    return STATUS_COLOUR.get(status)


def _status_fill(status_text: str) -> str | None:
    """Return a light dashboard fill without deriving or changing governance status."""
    status = status_text.strip().upper()
    if status == "COMPLETE":
        return HEX_STATUS_GREEN_LIGHT
    if status == "PENDING":
        return HEX_STATUS_AMBER_LIGHT
    if status in {"NOT APPLICABLE", "RISK ACCEPTED", "DRAFT"}:
        return HEX_STATUS_GREY_LIGHT
    return None


def extract_project_name(blocks: list[Block]) -> str:
    """Extract project name from first meaningful heading."""
    skip_patterns = {"stage 1", "proof of value", "dslc", "draft", "stage 1 readiness"}
    for b in blocks[:20]:
        if b.kind == "heading" and b.level <= 2:
            clean = _strip_md(b.text)
            if not any(p in clean.lower() for p in skip_patterns):
                return clean
    return "Project"


def _find_section(
    blocks: list[Block],
    *keywords: str,
    parent_range: tuple[int, int] | None = None,
) -> tuple[int, int]:
    """
    Find the start/end block indices of a section whose heading contains
    ALL given keywords (case-insensitive substring match).
    Searches within parent_range when supplied.
    Returns (-1, -1) if not found.
    """
    lo, hi = (0, len(blocks)) if parent_range is None else parent_range
    for idx in range(lo, hi):
        b = blocks[idx]
        if b.kind != "heading":
            continue
        h = b.text.lower()
        if all(kw.lower() in h for kw in keywords):
            level = b.level
            end   = hi
            for j in range(idx + 1, hi):
                if blocks[j].kind == "heading" and blocks[j].level <= level:
                    end = j
                    break
            return idx, end
    return -1, -1


def _load_logo(project_root: Path) -> Path | None:
    """
    Locate assets/sky_logo.png relative to the project root.
    Returns the path when found; prints a warning and returns None otherwise.
    The source image is never modified or resampled.
    """
    logo_path = project_root / "assets" / "sky_logo.png"
    if logo_path.exists():
        return logo_path
    print(f"WARNING: Sky logo not found at {logo_path}; document generated without logo.")
    return None


def _insert_logo_run(para, logo_path: Path, width_cm: float) -> bool:
    """
    Insert a logo picture into an existing paragraph run.
    Sized by width ONLY to preserve the original aspect ratio — never stretched.
    Returns True on success.
    """
    try:
        run = para.add_run()
        run.add_picture(str(logo_path), width=Cm(width_cm))
        return True
    except Exception as exc:
        print(f"WARNING: Could not insert logo: {exc}")
        return False


def extract_overall_readiness(blocks: list[Block]) -> str:
    """
    Extract Overall Readiness status verbatim from the source Markdown.
    Locates: Stage 1 Readiness → Overall Readiness.
    NEVER calculates or derives the status — returns the source value only.
    """
    s, e = _find_section(blocks, "stage 1 readiness")
    if s == -1:
        s, e = _find_section(blocks, "stage", "readiness")
    if s == -1:
        return ""
    os_, oe = _find_section(blocks, "overall readiness", parent_range=(s, e))
    if os_ == -1:
        os_, oe = _find_section(blocks, "overall", parent_range=(s, e))
    if os_ == -1:
        return ""
    for b in blocks[os_:oe]:
        if b.kind == "paragraph" and b.text.strip():
            return _strip_md(b.text).strip()
        if b.kind == "table" and b.rows:
            for row in b.rows:
                if len(row) >= 2 and row[1].strip():
                    return _strip_md(row[1]).strip()
    return ""


def extract_priority_blockers(blocks: list[Block]) -> list[str]:
    """
    Extract Priority Blockers from: Stage 1 Readiness → Priority Blockers.
    Returns up to 5 items from the source only.
    Never derives blockers from PENDING counts or scans other sections.
    """
    s, e = _find_section(blocks, "stage 1 readiness")
    if s == -1:
        s, e = _find_section(blocks, "stage", "readiness")
    if s == -1:
        print("WARNING: 'Stage 1 Readiness' not found; Priority Blockers omitted.")
        return []
    bs, be = _find_section(blocks, "priority blocker", parent_range=(s, e))
    if bs == -1:
        bs, be = _find_section(blocks, "blocker", parent_range=(s, e))
    if bs == -1:
        print("WARNING: 'Priority Blockers' subsection not found; omitted from cover page.")
        return []
    blockers: list[str] = []
    for b in blocks[bs:be]:
        if b.kind == "list":
            blockers.extend(_strip_md(item) for item in b.items if item.strip())
        elif b.kind == "table":
            for row in b.rows:
                if row and row[0].strip():
                    blockers.append(_strip_md(row[0]))
        if len(blockers) >= 5:
            break
    return blockers[:5]


def extract_project_summary(blocks: list[Block]) -> str:
    """
    Extract Project Summary from: Stage 1 Readiness → Project Summary.
    Falls back to Executive Summary (backward compatibility).
    Returns "" if neither found — caller omits the block and warns.
    """
    s, e = _find_section(blocks, "stage 1 readiness")
    if s == -1:
        s, e = _find_section(blocks, "stage", "readiness")
    if s != -1:
        ps, pe = _find_section(blocks, "project summary", parent_range=(s, e))
        if ps == -1:
            ps, pe = _find_section(blocks, "summary", parent_range=(s, e))
        if ps != -1:
            for b in blocks[ps:pe]:
                if b.kind == "paragraph" and len(b.text) > 40:
                    return _strip_md(b.text)[:500]
    # Backward-compatible fallback
    es, ee = _find_section(blocks, "executive summary")
    if es != -1:
        for b in blocks[es:ee]:
            if b.kind == "paragraph" and len(b.text) > 60:
                return _strip_md(b.text)[:500]
    print("WARNING: No 'Project Summary' or 'Executive Summary' found; omitted from cover page.")
    return ""


def extract_readiness_summary(blocks: list[Block]) -> list[list[str]]:
    """
    Extract the executive Readiness Summary from:
    Stage 1 Readiness → Readiness Summary.

    Preferred structure:
        Stage 1 Area | Status | Key Gap / Action | Go to Section

    Backward compatibility:
    - A legacy 3-column Area | Status | Summary table is accepted.
    - For legacy rows only, the third column is treated as Key Gap / Action
      and a stable section-name destination is added from the Stage 1 area.
    - No governance status or project fact is calculated here.
    """
    s, e = _find_section(blocks, "stage 1 readiness")
    if s == -1:
        s, e = _find_section(blocks, "stage", "readiness")
    if s == -1:
        print("WARNING: 'Stage 1 Readiness' section not found; Readiness Summary omitted.")
        return []

    rs, re_ = _find_section(blocks, "readiness summary", parent_range=(s, e))
    if rs == -1:
        rs, re_ = _find_section(blocks, "summary", parent_range=(s, e))
    if rs == -1:
        print("WARNING: 'Readiness Summary' subsection not found; omitted from cover page.")
        return []

    navigation_by_area = {
        "business": "Business & Use Case",
        "use case": "Business & Use Case",
        "data readiness": "Data Readiness",
        "exploratory": "Exploratory Analysis",
        "eda": "Exploratory Analysis",
        "technical development": "Feature Engineering / Model Development & Validation",
        "feature": "Feature Engineering",
        "model": "Model Development & Validation",
        "validation": "Model Development & Validation",
        "governance": "Governance & Sign-off",
        "final sign-off": "Governance & Sign-off",
        "sign-off": "Governance & Sign-off",
    }

    for b in blocks[rs:re_]:
        if b.kind != "table" or not b.rows:
            continue

        normalised_rows: list[list[str]] = []
        for row in b.rows[:8]:
            clean = [_strip_md(c) for c in row]

            if len(clean) >= 4:
                normalised_rows.append(clean[:4])
                continue

            if len(clean) == 3:
                area, status, gap = clean
                area_lower = area.lower()
                destination = ""
                for key, section_name in navigation_by_area.items():
                    if key in area_lower:
                        destination = section_name
                        break
                normalised_rows.append([area, status, gap, destination])
                continue

            if len(clean) >= 2:
                normalised_rows.append(clean + [""] * (4 - len(clean)))

        if normalised_rows:
            return normalised_rows

    print("WARNING: Readiness Summary table not found in expected location; omitted.")
    return []


def extract_governance_section_rows(blocks: list[Block]) -> list[list[str]]:
    """
    Extract governance table rows ONLY from the 'Governance & Sign-off' section.
    Never scans the whole document for governance keywords.
    Never extracts from Stage 1 Readiness, Risks & Limitations, or appendices.
    """
    s, e = _find_section(blocks, "governance", "sign")
    if s == -1:
        s, e = _find_section(blocks, "governance")
    if s == -1:
        print("WARNING: 'Governance & Sign-off' section not found.")
        return []
    for b in blocks[s:e]:
        if b.kind != "table":
            continue
        all_rows = ([b.headers] + b.rows) if b.headers else b.rows
        if not all_rows:
            continue
        first_lower = [c.lower() for c in all_rows[0]]
        if any("requirement" in c or "governance" in c or "sign" in c for c in first_lower):
            rows: list[list[str]] = []
            for row in b.rows:
                if not row:
                    continue
                req      = _strip_md(row[0]) if len(row) > 0 else ""
                status   = _strip_md(row[1]) if len(row) > 1 else ""
                evidence = _strip_md(row[2]) if len(row) > 2 else ""
                action   = _strip_md(row[3]) if len(row) > 3 else ""
                if req:
                    rows.append([req, status, evidence, action])
            if rows:
                return rows
    return []


def extract_governance_rows(blocks: list[Block]) -> list[list[str]]:
    """Governance rows from the authoritative section only (delegates to extract_governance_section_rows)."""
    return extract_governance_section_rows(blocks)


def extract_document_status(blocks: list[Block]) -> str:
    """
    Extract the document status from the source Markdown.
    Looks for explicit status fields near the top of the document
    (e.g. 'Status: DRAFT', 'Document Status: READY FOR REVIEW').
    Returns the source value verbatim, or 'DRAFT' if none is found.
    Never calculates or infers the status.
    """
    status_re = re.compile(
        r"(?:document\s+)?status\s*[:\|]\s*([A-Z][A-Z0-9 \-]+)",
        re.IGNORECASE,
    )
    # Only scan the first 30 blocks — status belongs near the top
    for b in blocks[:30]:
        if b.kind not in ("paragraph", "table"):
            continue
        candidates = [b.text] if b.kind == "paragraph" else [
            " | ".join(row) for row in ([b.headers] + b.rows if b.headers else b.rows)
        ]
        for text in candidates:
            m = status_re.search(text)
            if m:
                return m.group(1).strip()
    return "DRAFT"


def extract_dashboard_data(blocks: list[Block]) -> dict:
    """
    Assemble all executive cover page data from the source document.
    All content comes from the source — nothing is invented or calculated.
    """
    return {
        "project_name":      extract_project_name(blocks),
        "overall_readiness": extract_overall_readiness(blocks),
        "summary":           extract_project_summary(blocks),
        "readiness_rows":    extract_readiness_summary(blocks),
        "blockers":          extract_priority_blockers(blocks),
        "date":              date.today().strftime("%d %B %Y"),
    }


def split_for_appendices(
    blocks: list[Block],
) -> tuple[list[Block], dict[str, list[Block]]]:
    """
    Separate appendix sections from the main document body.

    - First occurrence of each appendix wins; later duplicates are ignored.
    - Output order is always A → B → C regardless of source order.
    - Never relies on dict insertion order alone.
    """
    appendix_triggers = {
        "evidence register":    "Appendix A",
        "data sources":         "Appendix B",
        "model technical detail": "Appendix C",
    }

    # First-occurrence mapping: label → (start_idx, end_idx)
    first_occurrence: dict[str, tuple[int, int]] = {}
    for idx, b in enumerate(blocks):
        if b.kind != "heading":
            continue
        heading_lower = b.text.lower()
        for trigger, label in appendix_triggers.items():
            if trigger in heading_lower:
                if label in first_occurrence:
                    print(f"WARNING: Duplicate appendix '{label}' detected; using first occurrence.")
                else:
                    level = b.level
                    end   = len(blocks)
                    for j in range(idx + 1, len(blocks)):
                        if blocks[j].kind == "heading" and blocks[j].level <= level:
                            end = j
                            break
                    first_occurrence[label] = (idx, end)
                break

    if not first_occurrence:
        return blocks, {}

    appendix_indices: set[int] = set()
    for s, e in first_occurrence.values():
        appendix_indices.update(range(s, e))

    main_blocks = [b for i, b in enumerate(blocks) if i not in appendix_indices]

    # Enforce canonical A → B → C output order
    canonical = ["Appendix A", "Appendix B", "Appendix C"]
    appendices: dict[str, list[Block]] = {
        label: blocks[first_occurrence[label][0]:first_occurrence[label][1]]
        for label in canonical
        if label in first_occurrence
    }
    return main_blocks, appendices


# ============================================================================
# DOCX HELPERS
# ============================================================================

def _shade_cell(cell, hex_colour: str) -> None:
    """Apply background shading to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_colour)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)


def _set_cell_margins(cell, top=60, bottom=60, left=100, right=100) -> None:
    """Set table cell padding."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _para_border_bottom(para, colour_hex: str = "0D3B72", sz: int = 6) -> None:
    """Add a bottom border to a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), colour_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_hyperlink(para, label: str, url: str, base_pt: float | None = None) -> None:
    """Add a clickable external hyperlink to a paragraph."""
    # Only allow explicit http(s) links. Other schemes remain plain text.
    if not re.match(r"^https?://", url.strip(), re.IGNORECASE):
        r = para.add_run(label)
        if base_pt:
            r.font.size = Pt(base_pt)
        return

    part = para.part
    r_id = part.relate_to(
        url.strip(),
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), HEX_SKY_BLUE)
    rPr.append(colour)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    if base_pt:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(base_pt * 2)))
        rPr.append(sz)

    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = label
    new_run.append(text_el)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


def _section_anchor(text: str) -> str:
    """Create a stable Word bookmark name from a report section heading or dashboard target."""
    text = _strip_md(text).strip()
    text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
    text = text.split("/")[0].strip()
    clean = re.sub(r"[^A-Za-z0-9_]", "_", text)
    clean = re.sub(r"_+", "_", clean).strip("_") or "section"
    if clean[0].isdigit():
        clean = f"section_{clean}"
    return clean[:40]


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """Add a Word bookmark to a paragraph."""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_hyperlink(para, label: str, anchor: str, base_pt: float | None = None) -> None:
    """Add a clickable hyperlink to a bookmark inside the same Word document."""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), HEX_SKY_BLUE)
    rPr.append(colour)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if base_pt:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(base_pt * 2)))
        rPr.append(sz)
    new_run.append(rPr)
    text_el = OxmlElement("w:t")
    text_el.text = label
    new_run.append(text_el)
    hyperlink.append(new_run)
    para._p.append(hyperlink)


def add_inline(para, text: str, base_pt: float | None = None) -> None:
    """Parse Markdown links, bold, italic and code into Word runs."""
    pattern = re.compile(
        r"(\[([^\]]+)\]\((https?://[^)]+)\)|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)"
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = para.add_run(text[pos:m.start()])
            if base_pt:
                r.font.size = Pt(base_pt)

        raw = m.group(0)
        if raw.startswith("["):
            _add_hyperlink(para, m.group(2), m.group(3), base_pt=base_pt)
        elif raw.startswith("**"):
            r = para.add_run(m.group(4))
            r.bold = True
            if base_pt:
                r.font.size = Pt(base_pt)
        elif raw.startswith("*"):
            r = para.add_run(m.group(5))
            r.italic = True
            if base_pt:
                r.font.size = Pt(base_pt)
        else:
            r = para.add_run(m.group(6))
            r.font.name = "Courier New"
            r.font.size = Pt(8)
        pos = m.end()

    if pos < len(text):
        r = para.add_run(text[pos:])
        if base_pt:
            r.font.size = Pt(base_pt)


# ============================================================================
# DOCUMENT SETUP
# ============================================================================

def setup_document(corporate_template: Path | None = None) -> Document:
    """
    Create and configure document.
    If corporate_template exists, load it as base; otherwise create from scratch.
    """
    if corporate_template and corporate_template.exists():
        doc = Document(str(corporate_template))
    else:
        doc = Document()

    # Page layout: A4
    for section in doc.sections:
        section.page_width  = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = section.right_margin = Cm(2.2)
        section.top_margin  = section.bottom_margin = Cm(2.0)

    # Body text style
    try:
        style = doc.styles["Normal"]
        style.font.name  = "Arial"
        style.font.size  = Pt(10.5)
        style.font.color.rgb = TEXT_DARK
        style.paragraph_format.space_after  = Pt(6)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.line_spacing = 1.15
    except:
        pass

    # Heading styles — Sky Blue H1, Sky Purple H2, charcoal H3/H4
    heading_defs = {
        "Heading 1": (16,   SKY_BLUE,   True),
        "Heading 2": (13,   SKY_PURPLE, True),
        "Heading 3": (11.5, TEXT_DARK,  True),
        "Heading 4": (10.5, TEXT_DARK,  False),
    }
    for name, (size, colour, bold) in heading_defs.items():
        try:
            hs = doc.styles[name]
            hs.font.name   = "Arial"
            hs.font.size   = Pt(size)
            hs.font.color.rgb = colour
            hs.font.bold   = bold
            hs.paragraph_format.space_before = Pt(12)
            hs.paragraph_format.space_after  = Pt(6)
            hs.paragraph_format.keep_with_next = True
        except:
            pass

    return doc


def add_header_footer(
    doc: Document,
    project_name: str,
    logo_path: Path | None = None,
    doc_status: str = "DRAFT",
) -> None:
    """Add Sky-branded header and footer to all sections."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # Normal-page header: logo left (if available), project name right
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.clear()

        if logo_path:
            hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _insert_logo_run(hp, logo_path, width_cm=1.8)
            hp.add_run("\t\t")  # tab to right position
        else:
            hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        r = hp.add_run(f"{project_name} | Stage 1 – Proof of Value | {doc_status}")
        r.font.name  = "Arial"
        r.font.size  = Pt(7.5)
        r.font.color.rgb = TEXT_GREY
        r.italic = True
        hp.paragraph_format.space_after = Pt(0)

        # Footer: Stage 1 | <doc_status> | Page X of Y
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        stage_r = fp.add_run(f"Stage 1 – Proof of Value  |  {doc_status}     Page ")
        stage_r.font.name  = "Arial"
        stage_r.font.size  = Pt(8)
        stage_r.font.color.rgb = TEXT_GREY

        run = fp.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")
        run._r.append(fldChar1)
        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = "PAGE"
        run._r.append(instrText)
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar2)
        run.font.size = Pt(8)
        run.font.color.rgb = TEXT_GREY

        of_r = fp.add_run(" of ")
        of_r.font.name = "Arial"
        of_r.font.size = Pt(8)
        of_r.font.color.rgb = TEXT_GREY

        run2 = fp.add_run()
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")
        run2._r.append(fldChar3)
        instrText2 = OxmlElement("w:instrText")
        instrText2.set(qn("xml:space"), "preserve")
        instrText2.text = "NUMPAGES"
        run2._r.append(instrText2)
        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")
        run2._r.append(fldChar4)
        run2.font.size = Pt(8)
        run2.font.color.rgb = TEXT_GREY


# ============================================================================
# EXECUTIVE READINESS PAGE
# ============================================================================

def _add_sky_divider(doc: Document, colour_hex: str = HEX_SKY_BLUE) -> None:
    """Insert a thin horizontal rule paragraph using a bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    _para_border_bottom(p, colour_hex, sz=4)


def add_cover_page(
    doc: Document,
    dash: dict,
    logo_path: Path | None = None,
) -> None:
    """
    Render the executive readiness cover page.

    Layout:
    Logo → Title → Project name → metadata →
    Overall Readiness → Readiness Summary →
    Priority Blockers → Project Summary → page break.

    The Readiness Summary is a reusable reviewer dashboard:
        Stage 1 Area | Status | Key Gap / Action | Go to Section

    All project content is drawn from the source Markdown.
    The renderer controls presentation only and never recalculates readiness.
    """
    # Sky logo
    if logo_path:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(0)
        lp.paragraph_format.space_after = Pt(14)
        _insert_logo_run(lp, logo_path, width_cm=3.0)

    # Title
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(0)
    tp.paragraph_format.space_after = Pt(4)
    tr = tp.add_run("Stage 1 – Proof of Value")
    tr.font.name = "Arial"
    tr.font.size = Pt(26)
    tr.font.color.rgb = SKY_BLUE
    tr.font.bold = True

    # Project name
    pp = doc.add_paragraph()
    pp.paragraph_format.space_before = Pt(0)
    pp.paragraph_format.space_after = Pt(8)
    pr = pp.add_run(dash["project_name"])
    pr.font.name = "Arial"
    pr.font.size = Pt(16)
    pr.font.color.rgb = TEXT_DARK

    # Document status + date
    mp = doc.add_paragraph()
    mp.paragraph_format.space_after = Pt(10)
    mr = mp.add_run(f"DRAFT — Human review required     |     {dash['date']}")
    mr.font.name = "Arial"
    mr.font.size = Pt(9)
    mr.font.color.rgb = TEXT_GREY
    mr.italic = True

    _add_sky_divider(doc, HEX_SKY_BLUE)

    # Overall Readiness — source value only.
    if dash.get("overall_readiness"):
        orh = doc.add_paragraph()
        orh.paragraph_format.space_before = Pt(12)
        orh.paragraph_format.space_after = Pt(3)
        ohr = orh.add_run("OVERALL READINESS")
        ohr.font.name = "Arial"
        ohr.font.size = Pt(8)
        ohr.font.color.rgb = TEXT_GREY
        ohr.font.bold = True

        ovp = doc.add_paragraph()
        ovp.paragraph_format.space_before = Pt(0)
        ovp.paragraph_format.space_after = Pt(8)
        ovr = ovp.add_run(dash["overall_readiness"])
        ovr.font.name = "Arial"
        ovr.font.size = Pt(18)
        ovr.font.bold = True

        # Important: negative phrases are checked first because
        # "NOT READY" contains the word "READY".
        status_lower = dash["overall_readiness"].lower()
        if any(kw in status_lower for kw in ("not ready", "blocking", "fail")):
            ovr.font.color.rgb = STATUS_RED
        elif any(kw in status_lower for kw in ("pending", "partial", "to confirm")):
            ovr.font.color.rgb = STATUS_AMBER
        elif any(kw in status_lower for kw in ("ready", "complete", "pass")):
            ovr.font.color.rgb = STATUS_GREEN
        else:
            ovr.font.color.rgb = SKY_PURPLE

    _add_sky_divider(doc, HEX_SKY_BLUE)

    # Readiness Summary — executive reviewer dashboard.
    if dash.get("readiness_rows"):
        rsh = doc.add_paragraph()
        rsh.paragraph_format.space_before = Pt(10)
        rsh.paragraph_format.space_after = Pt(6)
        rshr = rsh.add_run("READINESS SUMMARY")
        rshr.font.name = "Arial"
        rshr.font.size = Pt(8)
        rshr.font.color.rgb = TEXT_GREY
        rshr.font.bold = True

        headers = ["Stage 1 Area", "Status", "Key Gap / Action", "Go to Section"]
        tbl = doc.add_table(rows=1 + len(dash["readiness_rows"]), cols=4)
        tbl.style = "Table Grid"
        tbl.autofit = False

        # Header row
        for ci, h in enumerate(headers):
            cell = tbl.cell(0, ci)
            _shade_cell(cell, HEX_SKY_BLUE)
            _set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
            r.font.color.rgb = WHITE
            r.bold = True

        # Data rows
        for ri, row_data in enumerate(dash["readiness_rows"], start=1):
            row_data = list(row_data[:4]) + [""] * max(0, 4 - len(row_data))

            # Light alternate shading first.
            if ri % 2 == 0:
                for ci in range(4):
                    _shade_cell(tbl.cell(ri, ci), HEX_VERY_LIGHT_GREY)

            for ci in range(4):
                cell = tbl.cell(ri, ci)
                _set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = cell.paragraphs[0]
                val = _strip_md(row_data[ci])

                if ci == 1:
                    # Status: strong text + subtle traffic-light fill.
                    fill = _status_fill(val)
                    if fill:
                        _shade_cell(cell, fill)
                    sr = p.add_run(val)
                    sr.font.name = "Arial"
                    sr.font.size = Pt(8.5)
                    sr.bold = True
                    col = _status_colour(val)
                    sr.font.color.rgb = col[0] if col else TEXT_DARK

                elif ci == 3:
                    # Clickable navigation to a verified section bookmark in this document.
                    if val:
                        _add_internal_hyperlink(p, val, _section_anchor(val), base_pt=8)

                else:
                    add_inline(p, val, base_pt=8.5)

        # Widths sum to BODY_CM (16.6 cm).
        # Keep "Key Gap / Action" largest; navigation is intentionally compact.
        widths_cm = [3.5, 2.6, 7.1, 3.4]
        for row in tbl.rows:
            for ci, width in enumerate(widths_cm):
                row.cells[ci].width = Cm(width)

        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note.paragraph_format.space_after = Pt(2)
        nr = note.add_run(
            "Status text is authoritative; colour is a visual aid. "
            "See the named section for the detailed assessment and Appendix A for evidence."
        )
        nr.font.name = "Arial"
        nr.font.size = Pt(7.5)
        nr.font.color.rgb = TEXT_GREY
        nr.italic = True

    _add_sky_divider(doc, HEX_SKY_BLUE)

    # Priority Blockers
    if dash.get("blockers"):
        bh = doc.add_paragraph()
        bh.paragraph_format.space_before = Pt(10)
        bh.paragraph_format.space_after = Pt(6)
        bhr = bh.add_run("PRIORITY BLOCKERS")
        bhr.font.name = "Arial"
        bhr.font.size = Pt(8)
        bhr.font.color.rgb = TEXT_GREY
        bhr.font.bold = True

        for item in dash["blockers"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_before = Pt(1)
            bp.paragraph_format.space_after = Pt(1)
            br = bp.add_run(_strip_md(item))
            br.font.name = "Arial"
            br.font.size = Pt(10)
            br.font.color.rgb = TEXT_DARK

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    _add_sky_divider(doc, HEX_SKY_BLUE)

    # Project Summary
    if dash.get("summary"):
        psh = doc.add_paragraph()
        psh.paragraph_format.space_before = Pt(10)
        psh.paragraph_format.space_after = Pt(6)
        pshr = psh.add_run("PROJECT SUMMARY")
        pshr.font.name = "Arial"
        pshr.font.size = Pt(8)
        pshr.font.color.rgb = TEXT_GREY
        pshr.font.bold = True

        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        add_inline(sp, dash["summary"], base_pt=10.5)

    doc.add_page_break()



def add_table_of_contents(doc: Document) -> None:
    """Insert a clickable Word TOC and request a field refresh when Word opens it."""
    # Deliberately not a Heading style, so the TOC cannot include itself.
    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(10)
    hr = heading.add_run("Table of Contents")
    hr.font.name = "Arial"
    hr.font.size = Pt(18)
    hr.font.color.rgb = SKY_BLUE
    hr.bold = True
    _para_border_bottom(heading, HEX_SKY_BLUE, sz=4)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    run._r.append(instr)

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_sep)

    placeholder = p.add_run("Table of Contents — page numbers will update in Microsoft Word.")
    placeholder.font.name = "Arial"
    placeholder.font.size = Pt(9)
    placeholder.font.color.rgb = TEXT_GREY
    placeholder.italic = True

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    doc.add_page_break()


# ============================================================================
# GOVERNANCE TABLE
# ============================================================================

def add_governance_table(doc: Document, gov_rows: list[list[str]]) -> None:
    """Render the authoritative governance table with Sky theme."""
    if not gov_rows:
        return

    tbl = doc.add_table(rows=1 + len(gov_rows), cols=4)
    tbl.style = "Table Grid"

    headers = ["Governance Requirement", "Status", "Evidence / Gap", "Required Action"]
    for ci, h in enumerate(headers):
        cell = tbl.cell(0, ci)
        _shade_cell(cell, HEX_SKY_BLUE)
        _set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.color.rgb = WHITE
        r.bold = True

    for ri, row_data in enumerate(gov_rows, start=1):
        if ri % 2 == 0:
            for ci in range(4):
                _shade_cell(tbl.cell(ri, ci), HEX_VERY_LIGHT_GREY)
        for ci in range(4):
            cell = tbl.cell(ri, ci)
            _set_cell_margins(cell)
            val = row_data[ci] if ci < len(row_data) else ""
            p = cell.paragraphs[0]
            if ci == 1:  # Status column — apply status colour as text
                r = p.add_run(val)
                r.font.name = "Arial"
                r.font.size = Pt(8.5)
                r.bold = True
                col = _status_colour(val)
                r.font.color.rgb = col[0] if col else TEXT_DARK
            else:
                add_inline(p, val, base_pt=8.5)

    widths_cm = [4.5, 2.8, 4.8, 4.5]
    for row in tbl.rows:
        for ci, w in enumerate(widths_cm):
            if ci < len(row.cells):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ============================================================================
# BLOCK RENDERER
# ============================================================================

def _table_col_widths(headers: list[str]) -> list[float]:
    """Heuristic column widths in cm."""
    n = len(headers)
    if n == 1:
        return [BODY_CM]
    if n == 2:
        return [6.0, BODY_CM - 6.0]
    if n == 3:
        w = round(BODY_CM / 3, 1)
        return [w, w, BODY_CM - 2 * w]
    if n == 4:
        return [4.0, 2.5, 5.0, BODY_CM - 11.5]
    # ≥5 columns
    w = round(BODY_CM / n, 1)
    return [w] * n


def render_table(doc: Document, block: Block, small: bool = False) -> None:
    """Render a Markdown table block."""
    all_rows = [block.headers] + block.rows if block.headers else block.rows
    if not all_rows:
        return

    n_cols = max(len(r) for r in all_rows)
    tbl    = doc.add_table(rows=len(all_rows), cols=n_cols)
    tbl.style = "Table Grid"
    pt    = Pt(8) if small else Pt(9)
    widths = _table_col_widths(block.headers)

    for ri, row_data in enumerate(all_rows):
        cells = tbl.rows[ri].cells
        is_header = ri == 0 and bool(block.headers)

        if is_header:
            for c in cells:
                _shade_cell(c, HEX_SKY_BLUE)
        elif ri % 2 == 0:
            for c in cells:
                _shade_cell(c, HEX_VERY_LIGHT_GREY)

        for ci in range(n_cols):
            val = row_data[ci] if ci < len(row_data) else ""
            _set_cell_margins(cells[ci])
            p = cells[ci].paragraphs[0]
            if is_header:
                r = p.add_run(_strip_md(val))
                r.font.name = "Arial"
                r.font.size = pt
                r.font.color.rgb = WHITE
                r.bold = True
            else:
                add_inline(p, val, base_pt=float(pt.pt))

    # Apply widths
    for row in tbl.rows:
        for ci, w in enumerate(widths):
            if ci < len(row.cells):
                row.cells[ci].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def render_block(doc: Document, block: Block, in_appendix: bool = False, bookmark_counter: list[int] | None = None) -> None:
    """Render a single Block."""
    if block.kind == "heading":
        style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
        style_name = style_map.get(block.level, "Heading 4")
        p = doc.add_heading(level=block.level, text="")
        p.style = doc.styles[style_name]
        p.clear()
        add_inline(p, _strip_md(block.text))
        if not in_appendix and block.level <= 2 and bookmark_counter is not None:
            bookmark_counter[0] += 1
            _add_bookmark(p, _section_anchor(block.text), bookmark_counter[0])
        if not in_appendix:
            if block.level == 1:
                _para_border_bottom(p, HEX_SKY_BLUE, sz=4)
            elif block.level == 2:
                _para_border_bottom(p, HEX_SKY_PURPLE, sz=2)

    elif block.kind == "paragraph":
        p = doc.add_paragraph()
        add_inline(p, block.text, base_pt=10.5)
        p.paragraph_format.space_after = Pt(6)

    elif block.kind == "table":
        render_table(doc, block, small=in_appendix)

    elif block.kind == "list":
        list_style = "List Number" if block.ordered else "List Bullet"
        for item in block.items:
            p = doc.add_paragraph(style=list_style)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, item, base_pt=10)

    elif block.kind == "blockquote":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # Left border via XML
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        lft  = OxmlElement("w:left")
        lft.set(qn("w:val"),   "single")
        lft.set(qn("w:sz"),    "12")
        lft.set(qn("w:space"), "4")
        lft.set(qn("w:color"), HEX_SKY_ORANGE)
        pBdr.append(lft)
        pPr.append(pBdr)
        add_inline(p, block.text, base_pt=10)
        for run in p.runs:
            run.font.color.rgb = TEXT_GREY

    elif block.kind == "code":
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(block.text)
        r.font.name = "Courier New"
        r.font.size = Pt(8)
        r.font.color.rgb = TEXT_DARK

    elif block.kind == "rule":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _para_border_bottom(p, HEX_LIGHT_GREY, sz=4)


# ============================================================================
# APPENDICES
# ============================================================================

def add_appendices(doc: Document, appendices: dict[str, list[Block]]) -> None:
    """Render appendix sections."""
    if not appendices:
        return

    doc.add_page_break()
    ah = doc.add_heading("Appendices", level=1)
    _para_border_bottom(ah, HEX_SKY_BLUE, sz=4)

    for label, blocks in appendices.items():
        app_h = doc.add_heading(label, level=2)
        _para_border_bottom(app_h, HEX_SKY_PURPLE, sz=2)
        for b in blocks:
            if b.kind == "heading" and b.level <= 2:
                continue
            render_block(doc, b, in_appendix=True)

# ============================================================================
# PDF CONVERSION
# ============================================================================

def try_pdf(docx_path: Path) -> bool:
    """Attempt PDF conversion. Returns True on success."""
    try:
        from docx2pdf import convert
        pdf_path = docx_path.with_suffix(".pdf")
        convert(str(docx_path), str(pdf_path))
        return True
    except ImportError:
        pass
    except Exception:
        pass
    return False


# ============================================================================
# VALIDATION
# ============================================================================

def validate_before_render(md_path: Path) -> tuple[bool, str]:
    """
    Run validate_stage1.py against the Markdown file.
    Returns (passed, full_output_text).
    """
    validator = Path(__file__).parent / "validate_stage1.py"
    if not validator.exists():
        return True, "(validator not found — validation skipped)"
    try:
        result = subprocess.run(
            ["python", str(validator), str(md_path)],
            capture_output=True,
            text=True,
            timeout=60,
        )
        return result.returncode == 0, result.stdout
    except subprocess.TimeoutExpired:
        return False, "ERROR: Validation timed out"
    except Exception as exc:
        return False, f"ERROR: Validation failed: {exc}"


# ============================================================================
# MAIN
# ============================================================================

def _sanitise_filename(name: str) -> str:
    """Make filename safe."""
    return re.sub(r"[^a-zA-Z0-9_\- ]", "", name).strip().replace(" ", "_")


def render(
    md_path: Path,
    output_dir: Path,
    skip_validation: bool = False,
    force: bool = False,
) -> tuple[Path | None, Path | None, list[str]]:
    """
    Render Markdown to DOCX and optional PDF.
    Returns (docx_path_or_None, pdf_path_or_None, warnings_list).

    If validation fails and force=False, DOCX is NOT generated (docx_path=None).
    Use force=True only as a development override.
    """
    render_warnings: list[str] = []

    # Validation gate
    if not skip_validation:
        passed, val_output = validate_before_render(md_path)
        if not passed:
            if force:
                render_warnings.append("VALIDATION OVERRIDE USED — rendered despite ERROR findings")
                print(val_output)
            else:
                print(val_output)
                return None, None, ["Validation FAIL — use --force to override (development only)"]

    # Parse source
    lines  = md_path.read_text(encoding="utf-8").splitlines()
    blocks = parse_markdown(lines)

    # Locate Sky logo
    logo_path = _load_logo(Path(__file__).parent.parent)

    # Extract all dashboard data from source (nothing invented)
    dash      = extract_dashboard_data(blocks)
    gov_rows  = extract_governance_section_rows(blocks)
    doc_status = extract_document_status(blocks)

    # Split appendices: first occurrence wins, A→B→C order enforced
    main_blocks, appendices = split_for_appendices(blocks)

    # Identify the source governance table to skip (replaced by polished version)
    gov_table_idx = -1
    gs, ge = _find_section(main_blocks, "governance", "sign")
    if gs == -1:
        gs, ge = _find_section(main_blocks, "governance")
    if gs != -1:
        for i in range(gs, ge):
            if main_blocks[i].kind == "table":
                hdrs = [c.lower() for c in main_blocks[i].headers]
                if any("requirement" in c or "governance" in c or "sign" in c for c in hdrs):
                    gov_table_idx = i
                    break

    # Prepare output path
    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = _sanitise_filename(dash["project_name"])
    docx_path = output_dir / f"Stage1_{safe_name}.docx"

    # Corporate template support
    corporate_template = Path(__file__).parent.parent / "templates" / "corporate_stage1_template.docx"

    # Build DOCX
    doc = setup_document(corporate_template if corporate_template.exists() else None)
    add_header_footer(doc, dash["project_name"], logo_path=logo_path, doc_status=doc_status)
    add_cover_page(doc, dash, logo_path=logo_path)
    add_table_of_contents(doc)

    # Render main content
    # Source governance table is replaced by the polished version (one table only)
    gov_inserted = False
    bookmark_counter = [0]
    for i, b in enumerate(main_blocks):
        if i == gov_table_idx:
            continue  # skip source table — replaced below

        render_block(doc, b, bookmark_counter=bookmark_counter)

        # After governance heading, insert polished table exactly once
        if (not gov_inserted and gov_rows and b.kind == "heading"
                and ("governance" in b.text.lower() or "sign" in b.text.lower())):
            gov_inserted = True
            add_governance_table(doc, gov_rows)

    add_appendices(doc, appendices)

    # Save the completed editable Word document.
    doc.save(str(docx_path))

    # Attempt PDF conversion where locally supported.
    pdf_path = None
    if try_pdf(docx_path):
        pdf_path = docx_path.with_suffix(".pdf")

    return docx_path, pdf_path, render_warnings


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Render a validated Stage 1 Markdown to professional DOCX and optional PDF.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Normal workflow:
  python tools/render_stage1.py stage1_draft.md
  python tools/render_stage1.py stage1_draft.md --output-dir output

Development overrides (not for normal use):
  python tools/render_stage1.py stage1_draft.md --force
  python tools/render_stage1.py stage1_draft.md --skip-validation
        """,
    )
    parser.add_argument("markdown_file", help="Validated Stage 1 Markdown file")
    parser.add_argument("--output-dir", default="output", help="Output directory (default: output/)")
    parser.add_argument(
        "--force",
        action="store_true",
        help="Render despite validation ERROR findings (development override; not for normal use)",
    )
    parser.add_argument(
        "--skip-validation",
        action="store_true",
        help="Skip validation entirely (development/testing only)",
    )
    args = parser.parse_args()

    md_path = Path(args.markdown_file)
    if not md_path.exists():
        print(f"ERROR: File not found: {md_path}")
        sys.exit(1)

    output_dir = Path(args.output_dir)

    print("=" * 70)
    print("DSLC STAGE 1 DOCUMENT RENDERER")
    print("=" * 70)
    print()
    print(f"Input  : {md_path}")
    print()

    docx_path, pdf_path, render_warnings = render(
        md_path,
        output_dir,
        skip_validation=args.skip_validation,
        force=args.force,
    )

    if docx_path is None:
        print("=" * 70)
        print("RENDER FAILED")
        print("=" * 70)
        print()
        print("Validation returned ERROR findings. Resolve them before rendering.")
        print("Use --force to override (development only).")
        print()
        sys.exit(1)

    # Determine validation label for completion message
    if args.skip_validation:
        val_label = "SKIPPED"
    elif args.force:
        val_label = "OVERRIDDEN (--force used)"
    else:
        val_label = "PASS"

    print("=" * 70)
    print("DSLC STAGE 1 RENDER COMPLETE")
    print("=" * 70)
    print()

    if args.force and any("OVERRIDE" in w for w in render_warnings):
        print("VALIDATION OVERRIDE USED")
        print()

    print(f"Validation:  {val_label}")
    print()
    print(f"DOCX:")
    print(f"  {docx_path}")
    print()
    if pdf_path:
        print(f"PDF:")
        print(f"  {pdf_path}")
    else:
        print("PDF:  unavailable — DOCX is the editable master")
    print()
    warning_count = len(render_warnings)
    print(f"Warnings:  {warning_count}")
    if render_warnings:
        for w in render_warnings:
            print(f"  {w}")
    print("=" * 70)
    sys.exit(0)


if __name__ == "__main__":
    main()